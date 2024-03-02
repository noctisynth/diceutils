from typing import Dict, List, Any, Optional
from reportlab.pdfgen import canvas
from reportlab.platypus import Paragraph, SimpleDocTemplate, Image as PDFImage, Table
from docx.document import Document as DocumentType
from docx import Document
from enum import Enum

import abc
import re
import random


class Element(metaclass=abc.ABCMeta):
    def __init__(self, _type: str, content: str) -> None:
        self.type: str = _type
        self.content: str = content
        self._tag: str = "act"

    def __repr__(self) -> str:
        return f"Element(type={self.type}, content={self.content})"

    @property
    def tag(self) -> str:
        return self._tag


class Text(Element):
    def __init__(self, text, tag: Optional[str] = None):
        super().__init__("text", text)
        self._tag = tag or self._tag

    def __repr__(self) -> str:
        return f"Text(type={self.type!r}, content={self.content!r}, tag={self._tag!r})"

    def set_tag(self, tag: str):
        self._tag = tag


class Image(Element):
    def __init__(self, url) -> None:
        super().__init__("image", url)

    def __repr__(self) -> str:
        return f"Image(type={self.type!r}, url={self.content!r})"


class Role(Enum):
    DICER = -1
    GM = 0
    PL = 1
    OB = 2


class Message:
    def __init__(
        self,
        user_code: str,
        role: Role,
        nickname: str,
        date: str,
        elements: List[Element] = [],
    ):
        self.user_code = user_code
        self.nickname = nickname
        self.role = role
        self.date = date
        self.elements = elements

    def __repr__(self) -> str:
        return f"Message(user_code={self.user_code!r}, role={self.role}, nickname={self.nickname!r}, elements={self.elements!r}, date={self.date!r})"


class Messages(List[Message]):
    def __init__(self, messages: List[Message] = []):
        super().__init__(messages)

    def append(self, message: Message):
        return super().append(message)

    def add_message(
        self,
        user_code: str,
        role: Role,
        nickname: str,
        date: str,
        content: List[Dict[str, Any]],
    ):
        elements = []
        try:
            for ele in content:
                _type = ele["type"]
                data = ele["data"]
                if _type == "text":
                    text = data["text"]
                    elements.append(Text(text))
                elif _type == "image":
                    url = data["url"]
                    elements.append(Image(url))
        except KeyError as e:
            raise KeyError(f"Message format error. Missing key field: {e}")
        message = Message(user_code, role, nickname, date, elements)
        self.append(message)


class ExportConfig:
    def __init__(self):
        self.first_line_indent = True  # 首行缩进
        self.dice_command_filter = False  # 骰子指令过滤
        self.external_comment_filter = False  # 场外发言过滤
        self.display_image = True  # 图片显示
        self.display_datetime = True  # 时间显示
        self.display_account = True  # 帐号显示
        self.display_year_month_day = False  # 年月日显示


class Renderer(metaclass=abc.ABCMeta):
    @staticmethod
    def split_and_label(text: str) -> Dict[str, str]:
        pattern = r"“[^”]*”|[^“”]+"
        parts: List[str] = re.findall(pattern, text)

        result_dict = {}

        for part in parts:
            if part.startswith("“") and part.endswith("”"):
                result_dict[part] = "speak"
            else:
                result_dict[part] = "act"

        return result_dict

    @staticmethod
    def parse_message(message: Message, config: ExportConfig) -> Optional[Message]:
        elements = message.elements

        if len(elements) == 0:
            return None

        if config.dice_command_filter and message.role == -1:
            return None

        if len(elements) == 1 and isinstance(elements[0], Image):
            return message if config.display_image else None

        ele_iter = iter(elements)
        first_ele = next(ele_iter)
        is_external_comment = first_ele.content.startswith(("(", "（"))
        is_text = isinstance(first_ele, Text)
        if config.external_comment_filter and is_text and is_external_comment:
            return None

        if is_text and is_external_comment:
            first_ele.set_tag("outside")
            for ele in ele_iter:
                if isinstance(ele, Text):
                    ele.set_tag("outside")
            return message

        new_elements = []
        if is_text:
            for text, tag in Renderer.split_and_label(first_ele.content).items():
                new_elements.append(Text(text, tag))

        for ele in ele_iter:
            if isinstance(ele, Text):
                for text, tag in Renderer.split_and_label(ele.content).items():
                    new_elements.append(Text(text, tag))

        message.elements = new_elements
        return message

    @abc.abstractmethod
    def render_message(self, message: Message) -> None:
        raise NotImplementedError

    @staticmethod
    def render(
        messages: Messages,
        renderer: "Renderer",
        config: Optional[ExportConfig] = None,
    ) -> "Renderer":
        for message in messages:
            if message := renderer.parse_message(message, config or ExportConfig()):
                renderer.render_message(message)
        return renderer

    @abc.abstractmethod
    def export(self, filename: str) -> None:
        raise NotImplementedError


# class PDFRenderer(Renderer):
#     _canvas: canvas.Canvas

#     def __init__(self) -> None:
#         self._canvas = canvas.Canvas(str(random.randbytes(5)) + ".pdf")

#     def export(self, filename: str) -> None:
#         return self._canvas._doc.SaveToFile(filename=filename, canvas=self._canvas)

#     def render_message(self, message: Message):
#         self._canvas.drawText(Paragraph("dde"))
#         for element in message.elements:
#             if element.type == "text":
#                 if element.tag == "act":
#                     ...
#                 elif element.tag == "outside":
#                     ...
#                 elif element.tag == "speak":
#                     ...


# class DocxRenderer(Renderer):
#     document: DocumentType

#     def __init__(self) -> None:
#         self.document = Document()
#         self.document.styles["Normal"].font.name = "SimSun"

#     def export(self) -> None:
#         self.document.save("a.docx")

#     def render_message(self, message: Message) -> None:
#         paragraph = self.document.add_paragraph()
#         for element in message.elements:
#             if element.type == "text":
#                 if element.tag == "act":
#                     run = paragraph.add_run(element.content)
#                     run.bold = True
#                     run.font.name = "SimSun"
#                 elif element.tag == "outside":
#                     paragraph.add_run("（")
#                     content = element.content.strip("（）()")
#                     run = paragraph.add_run(content)
#                     run.font.color.rgb = (207, 210, 210)
#                     run.font.name = "SimSun"
#                 elif element.tag == "speak":
#                     content = element.content.strip('“”""')
#                     paragraph.add_run("“")
#                     run = paragraph.add_run(content)
#                     run.font.name = "Microsoft YaHei"
#                     paragraph.add_run("”")
#             else:
#                 # self.document.add_picture()
#                 ...

#         return
