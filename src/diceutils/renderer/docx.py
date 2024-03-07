from diceutils.renderer._models import Renderer, Message, Role, Element
from pathlib import Path
from typing import Optional

from docx.enum.style import WD_STYLE_TYPE
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

import docx


class DocxRenderer(Renderer):
    document: Document

    def __init__(self) -> None:
        self.document = docx.Document()
        self._add_chinese_font_style("微软雅黑")
        self._add_chinese_font_style("SimSun", "SimSun")
        self._add_chinese_font_style("楷体")

    def _add_chinese_font_style(
        self, style_name: str, font_name: Optional[str] = None, font_size: float = 12
    ):
        font_name = font_name or style_name
        style_cn = self.document.styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
        style_cn.font.name = font_name  # type: ignore
        style_cn.font.size = Pt(font_size)  # type: ignore
        style_cn.element.rPr.rFonts.set(qn("w:eastAsia"), font_name)

    def _render_keeper(self, element: Element, paragraph: Paragraph) -> None:
        paragraph.add_run(element.content, style="楷体")

    def _render_act(
        self, element: Element, paragraph: Paragraph, is_first: bool = False
    ) -> None:
        content = "#" + element.content.strip("#＃") if is_first else element.content
        run = paragraph.add_run(content, style="SimSun")
        run.font.italic = True

    def _render_dicer(self, element: Element, paragraph: Paragraph) -> None:
        paragraph.add_run(element.content, style="微软雅黑")

    def _render_command(self, element: Element, paragraph: Paragraph) -> None:
        run = paragraph.add_run(element.content, style="微软雅黑")
        run.font.name = "Cascadia Mono"

    def _render_speak(self, element: Element, paragraph: Paragraph) -> None:
        content = "“" + element.content.strip('"“”') + "”"
        paragraph.add_run(content, style="SimSun")

    def _render_outside(self, element: Element, paragraph: Paragraph) -> None:
        content = "（" + element.content.strip("()（）") + "）"
        run = paragraph.add_run(content, style="SimSun")
        run.font.color.rgb = RGBColor(207, 210, 210)

    def _render_image(self, image: Element, paragraph: Paragraph) -> None:
        run = paragraph.add_run(f"[IMAGE '{image.content}']")
        run.font.name = "Cascadia Mono"

    def render_message(self, message: Message) -> None:
        paragraph = self.document.add_paragraph()
        date = paragraph.add_run(message.date)
        date.font.name = "Cascadia Mono"
        date.font.color.rgb = RGBColor(207, 210, 210)
        paragraph.add_run("<" + message.nickname + ">", style="微软雅黑")
        for index, element in enumerate(message.elements):
            is_first = index == 0
            if element.type == "text":
                if message.role == Role.DICER:
                    self._render_dicer(element, paragraph)
                    continue
                elif message.role == Role.OB:
                    self._render_outside(element, paragraph)
                    continue

                if element.tag == "act":
                    if message.role == Role.GM:
                        self._render_keeper(element, paragraph)
                    else:
                        self._render_act(element, paragraph, is_first=is_first)
                elif element.tag == "outside":
                    self._render_outside(element, paragraph)
                elif element.tag == "speak":
                    self._render_speak(element, paragraph)
                elif element.tag == "command":
                    self._render_command(element, paragraph)
            elif element.type == "image":
                self._render_image(element, paragraph)

    def export(self, filename: str) -> Path:
        filepath = Path("./exports/" + filename + ".docx").resolve()
        filepath.parent.mkdir(parents=True, exist_ok=True)
        self.document.save(str(filepath))
        return filepath
